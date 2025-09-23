from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import re

def create_word_document():
    """Convert the solution document to Microsoft Word format with proper styling"""
    
    # Create a new document
    doc = Document()
    
    # Set up custom styles
    setup_document_styles(doc)
    
    # Read the markdown content
    with open('SOLUTION_DOCUMENT.md', 'r', encoding='utf-8') as file:
        content = file.read()
    
    # Process the content
    process_markdown_content(doc, content)
    
    # Save the document
    doc.save('AI_Contract_Leakage_Detection_Solution.docx')
    print("Word document created successfully: AI_Contract_Leakage_Detection_Solution.docx")

def setup_document_styles(doc):
    """Set up custom styles for the document"""
    
    # Title style
    title_style = doc.styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Calibri'
    title_font.size = Pt(24)
    title_font.bold = True
    title_font.color.rgb = None  # Default color
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(18)
    
    # Subtitle style
    subtitle_style = doc.styles.add_style('Custom Subtitle', WD_STYLE_TYPE.PARAGRAPH)
    subtitle_font = subtitle_style.font
    subtitle_font.name = 'Calibri'
    subtitle_font.size = Pt(16)
    subtitle_font.italic = True
    subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_style.paragraph_format.space_after = Pt(24)
    
    # Main heading style (H1)
    h1_style = doc.styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    h1_font = h1_style.font
    h1_font.name = 'Calibri'
    h1_font.size = Pt(20)
    h1_font.bold = True
    h1_style.paragraph_format.space_before = Pt(24)
    h1_style.paragraph_format.space_after = Pt(12)
    
    # Sub heading style (H2)
    h2_style = doc.styles.add_style('Custom Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    h2_font = h2_style.font
    h2_font.name = 'Calibri'
    h2_font.size = Pt(16)
    h2_font.bold = True
    h2_style.paragraph_format.space_before = Pt(18)
    h2_style.paragraph_format.space_after = Pt(10)
    
    # Sub-sub heading style (H3)
    h3_style = doc.styles.add_style('Custom Heading 3', WD_STYLE_TYPE.PARAGRAPH)
    h3_font = h3_style.font
    h3_font.name = 'Calibri'
    h3_font.size = Pt(14)
    h3_font.bold = True
    h3_style.paragraph_format.space_before = Pt(12)
    h3_style.paragraph_format.space_after = Pt(8)
    
    # Code block style
    code_style = doc.styles.add_style('Custom Code', WD_STYLE_TYPE.PARAGRAPH)
    code_font = code_style.font
    code_font.name = 'Consolas'
    code_font.size = Pt(10)
    code_style.paragraph_format.left_indent = Inches(0.5)
    code_style.paragraph_format.space_before = Pt(6)
    code_style.paragraph_format.space_after = Pt(6)

def process_markdown_content(doc, content):
    """Process markdown content and convert to Word format"""
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # Title (first line)
        if i == 0 and line.startswith('# '):
            title = line[2:].strip()
            p = doc.add_paragraph(title, style='Custom Title')
            i += 1
            continue
        
        # Subtitle (second line if it's a heading)
        if i == 1 and line.startswith('## '):
            subtitle = line[3:].strip()
            p = doc.add_paragraph(subtitle, style='Custom Subtitle')
            i += 1
            continue
        
        # Main headings (##)
        if line.startswith('## ') and not line.startswith('### '):
            heading = line[3:].strip()
            doc.add_paragraph(heading, style='Custom Heading 1')
            i += 1
            continue
        
        # Sub headings (###)
        if line.startswith('### ') and not line.startswith('#### '):
            heading = line[4:].strip()
            doc.add_paragraph(heading, style='Custom Heading 2')
            i += 1
            continue
        
        # Sub-sub headings (####)
        if line.startswith('#### '):
            heading = line[5:].strip()
            doc.add_paragraph(heading, style='Custom Heading 3')
            i += 1
            continue
        
        # Tables
        if '|' in line and '---' not in line:
            table_lines = []
            while i < len(lines) and ('|' in lines[i] or lines[i].strip() == ''):
                if lines[i].strip() and '---' not in lines[i]:
                    table_lines.append(lines[i])
                i += 1
            
            if table_lines:
                create_table(doc, table_lines)
            continue
        
        # Code blocks
        if line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            if code_lines:
                create_code_block(doc, code_lines)
            i += 1
            continue
        
        # Bullet points
        if line.startswith('- ') or line.startswith('* '):
            bullet_lines = []
            while i < len(lines) and (lines[i].strip().startswith('- ') or lines[i].strip().startswith('* ') or lines[i].strip() == ''):
                if lines[i].strip():
                    bullet_lines.append(lines[i].strip()[2:])
                i += 1
            
            for bullet in bullet_lines:
                p = doc.add_paragraph(bullet, style='List Bullet')
            continue
        
        # Regular paragraphs
        if line and not line.startswith('#') and not line.startswith('---'):
            # Handle bold text
            paragraph = doc.add_paragraph()
            add_formatted_text(paragraph, line)
        
        i += 1

def create_table(doc, table_lines):
    """Create a table from markdown table lines"""
    if not table_lines:
        return
    
    # Parse header
    headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
    
    # Parse data rows
    data_rows = []
    for line in table_lines[1:]:
        if '---' not in line:
            row = [cell.strip() for cell in line.split('|')[1:-1]]
            if row:
                data_rows.append(row)
    
    if not data_rows:
        return
    
    # Create table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Add headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        if i < len(header_cells):
            header_cells[i].text = header
            # Make header bold
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
    
    # Add data rows
    for row_data in data_rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            if i < len(row_cells):
                row_cells[i].text = cell_data

def create_code_block(doc, code_lines):
    """Create a code block with proper formatting"""
    for line in code_lines:
        doc.add_paragraph(line, style='Custom Code')

def add_formatted_text(paragraph, text):
    """Add text with formatting (bold, italic) to a paragraph"""
    # Simple bold formatting
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        else:
            # Regular text
            paragraph.add_run(part)

if __name__ == "__main__":
    create_word_document()
